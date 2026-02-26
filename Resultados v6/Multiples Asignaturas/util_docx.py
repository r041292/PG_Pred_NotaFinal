# util_docx.py
import io, textwrap
from contextlib import redirect_stdout, redirect_stderr
import matplotlib
matplotlib.use("Agg")  # si estás en headless
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

class captura_a_docx:
    def __init__(self, ruta_docx, titulo=None, ancho_imagen_pulg=6.5, ancho_texto_chars=110):
        self.ruta_docx = ruta_docx
        self.titulo = titulo
        self.ancho_imagen_pulg = ancho_imagen_pulg
        self.ancho_texto_chars = ancho_texto_chars

    def __enter__(self):
        # Documento listo desde el inicio para insertar en orden
        self.doc = Document()
        if self.titulo:
            self.doc.add_heading(self.titulo, level=0)

        # Captura de consola
        self._buf = io.StringIO()
        self._buf_pos = 0  # <-- puntero para saber qué parte ya volcamos
        self._out_cm = redirect_stdout(self._buf)
        self._err_cm = redirect_stderr(self._buf)
        self._out_cm.__enter__(); self._err_cm.__enter__()

        # Guardar figs iniciales para detectar nuevas si quisieras
        self._start_figs = set(plt.get_fignums())
        return self

    def _add_text(self, texto):
        if not texto.strip():
            return
        p = self.doc.add_paragraph()
        run = p.add_run("\n".join(textwrap.fill(l, self.ancho_texto_chars, replace_whitespace=False)
                                  for l in texto.splitlines()))
        run.font.name = "Consolas"
        # forzar fuente monospace en Word
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "Consolas")
        run.font.size = Pt(10)

    def flush_consola(self):
        """Inserta en el Word el texto NUEVO impreso desde el último flush."""
        contenido = self._buf.getvalue()
        nuevo = contenido[self._buf_pos:]
        if nuevo:
            self._add_text(nuevo)
            self._buf_pos = len(contenido)

    def insertar_figura_actual(self, fig=None, width_inches=None, salto_pagina=False):
        """Inserta inmediatamente la figura actual (en el punto actual del documento)."""
        if fig is None:
            fig = plt.gcf()
        if width_inches is None:
            width_inches = self.ancho_imagen_pulg

        bio = io.BytesIO()
        fig.savefig(bio, format="png", dpi=200, bbox_inches="tight")
        bio.seek(0)
        self.doc.add_picture(bio, width=Inches(width_inches))
        if salto_pagina:
            self.doc.add_page_break()

    def __exit__(self, exc_type, exc, tb):
        # cerrar redirecciones
        self._out_cm.__exit__(exc_type, exc, tb)
        self._err_cm.__exit__(exc_type, exc, tb)

        # volcar cualquier texto que haya quedado pendiente
        self.flush_consola()

        # guardar
        self.doc.save(self.ruta_docx)



'''

# === util_docx.py ===
import io, textwrap
from contextlib import redirect_stdout, redirect_stderr
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK

class captura_a_docx:
    """
    Requiere: pip install python-docx
    Usa:
    with captura_a_docx("salida.docx", titulo="Resultados"):
        # prints + figuras
    """
    def __init__(self, ruta_docx, titulo=None, ancho_imagen_pulg=6.5, ancho_texto_chars=110):
        self.ruta_docx = ruta_docx
        self.titulo = titulo
        self.ancho_imagen_pulg = ancho_imagen_pulg
        self.ancho_texto_chars = ancho_texto_chars
        self._figs_forzadas = []

    def __enter__(self):
        self._buf = io.StringIO()
        self._start_figs = set(plt.get_fignums())
        self._out_cm = redirect_stdout(self._buf)
        self._err_cm = redirect_stderr(self._buf)
        self._out_cm.__enter__(); self._err_cm.__enter__()
        return self
    
    def anexar_figura_actual(self, fig=None):
        import matplotlib.pyplot as plt
        if fig is None:
            fig = plt.gcf()
        if fig not in self._figs_forzadas:
            self._figs_forzadas.append(fig)

    def insertar_figura_actual(self, fig=None, width_inches=None, salto_pagina=False):
        """
        Inserta inmediatamente la figura actual en el documento, en el orden de ejecución.
        """
        import matplotlib.pyplot as plt
        import io
        from docx.shared import Inches
        from docx.enum.text import WD_BREAK

        if fig is None:
            fig = plt.gcf()
        if width_inches is None:
            width_inches = self.ancho_imagen_pulg

        # Crear objeto BytesIO y guardar figura como PNG
        bio = io.BytesIO()
        fig.savefig(bio, format="png", dpi=200, bbox_inches="tight")
        bio.seek(0)

        # Si el documento no existe aún, créalo
        if not hasattr(self, "doc"):
            from docx import Document
            self.doc = Document()
            if self.titulo:
                self.doc.add_heading(self.titulo, level=0)

        # Añadir la imagen en el punto actual
        self.doc.add_picture(bio, width=Inches(width_inches))
        if salto_pagina:
            self.doc.add_page_break()

    #nuevo exit         
    def __exit__(self, exc_type, exc, tb):
        self._out_cm.__exit__(exc_type, exc, tb)
        self._err_cm.__exit__(exc_type, exc, tb)

        texto = self._buf.getvalue()
        # Si ya existe self.doc (porque usaste insertar_figura_actual), la reusamos
        if not hasattr(self, "doc"):
            from docx import Document
            self.doc = Document()
            if self.titulo:
                self.doc.add_heading(self.titulo, level=0)

        if texto.strip():
            p = self.doc.add_paragraph()
            run = p.add_run(texto)
            run.font.name = "Consolas"

        self.doc.save(self.ruta_docx)

Exit original 
    def __exit__(self, exc_type, exc, tb):
        self._out_cm.__exit__(exc_type, exc, tb)
        self._err_cm.__exit__(exc_type, exc, tb)

        texto = self._buf.getvalue()
        nuevas_figs_nums = [n for n in plt.get_fignums() if n not in self._start_figs]
        nuevas_figs = [plt.figure(n) for n in nuevas_figs_nums]

        nuevas_figs_nums = [n for n in plt.get_fignums() if n not in self._start_figs]
        nuevas_figs = [plt.figure(n) for n in nuevas_figs_nums]

        # <-- Agrega las forzadas (sin duplicar)
        for f in self._figs_forzadas:
            if f not in nuevas_figs:
                nuevas_figs.append(f)

        doc = Document()
        if self.titulo:
            doc.add_heading(self.titulo, level=0)

        if texto.strip():
            p = doc.add_paragraph()
            # monospace “consolas” (si no está disponible, Word usa una alternativa)
            run = p.add_run("\n".join(textwrap.fill(l, self.ancho_texto_chars) for l in texto.splitlines()))
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Consolas")
            run.font.size = Pt(10)

        # Insertar cada figura como imagen
        for i, fig in enumerate(nuevas_figs, start=1):
            bio = io.BytesIO()
            fig.savefig(bio, format="png", dpi=200, bbox_inches="tight")
            bio.seek(0)
            doc.add_picture(bio, width=Inches(self.ancho_imagen_pulg))
            if i != len(nuevas_figs):
                doc.add_page_break()

        doc.save(self.ruta_docx)

        for f in nuevas_figs:
            plt.close(f)
 

'''
